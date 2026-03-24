"""
Gerador de Documentos Jurídicos
Daniel Cunha Detter — OAB/SP 258.095
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import os
import sys
import re
import subprocess
import shutil
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None


# ── Paleta ──────────────────────────────────────────────────────────────────
C = {
    "bg":        "#0e0e16",
    "bg2":       "#14141f",
    "bg3":       "#1c1c2e",
    "panel":     "#12121c",
    "border":    "#2a2a3e",
    "gold":      "#c9a84c",
    "gold_dim":  "#8a6f2e",
    "gold_light":"#e8d08a",
    "text":      "#e8e4dc",
    "muted":     "#7a7490",
    "accent":    "#3a3a5c",
    "success":   "#2d6a4f",
    "success_fg":"#6ee7b7",
    "error":     "#7f1d1d",
    "error_fg":  "#fca5a5",
    "white":     "#ffffff",
    "sep":       "#1e1e30",
}

FONT_TITLE  = ("Georgia", 22, "bold")
FONT_SUB    = ("Georgia", 12, "italic")
FONT_LABEL  = ("Georgia", 9, "bold")
FONT_INPUT  = ("Consolas", 11)
FONT_BTN    = ("Georgia", 11, "bold")
FONT_SMALL  = ("Consolas", 9)
FONT_STATUS = ("Consolas", 10)
FONT_SECT   = ("Georgia", 10, "bold")


def resource_path(rel):
    """Caminho correto dentro do .exe (PyInstaller) ou dev."""
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, rel)


# ── Lógica de geração ────────────────────────────────────────────────────────

def substituir_paragrafo(text, dados, rg_c, cpf_c, nt_c, rt_c):
    r = text
    simples = {
        "NOME DO MENOR/INCAPAZ":      dados.get("nome_menor", ""),
        "Data de Nascimento":          dados.get("nascimento", ""),
        "NOME DA REPRESENTANTE LEGAL": dados.get("nome_rep", ""),
        "estado civil":                dados.get("estado_civil", ""),
        "profissão":                   dados.get("profissao", ""),
        "Endereço completo com CEP":   dados.get("endereco", ""),
        "Cidade":                      dados.get("cidade", ""),
        "data":                        dados.get("data_dia", ""),
        "mês":                         dados.get("mes", ""),
        "ano":                         dados.get("ano", ""),
        "e-mail do advogado":          dados.get("email", ""),
    }
    for campo, val in simples.items():
        if val:
            r = r.replace(f"[{campo}]", val)

    rg_v = [dados.get("rg_menor",""), dados.get("rg_rep","")]
    def rg_r(m):
        i = rg_c[0]; rg_c[0] += 1
        return rg_v[i] if i < len(rg_v) and rg_v[i] else m.group(0)
    r = re.sub(r'\[nº do RG\]', rg_r, r)

    cpf_v = [dados.get("cpf_menor",""), dados.get("cpf_rep","")]
    def cpf_r(m):
        i = cpf_c[0]; cpf_c[0] += 1
        return cpf_v[i] if i < len(cpf_v) and cpf_v[i] else m.group(0)
    r = re.sub(r'\[nº do CPF\]', cpf_r, r)

    nt_v = [dados.get("test1_nome",""), dados.get("test2_nome","")]
    def nt_r(m):
        i = nt_c[0]; nt_c[0] += 1
        return nt_v[i] if i < len(nt_v) and nt_v[i] else m.group(0)
    r = re.sub(r'\[NOME\]', nt_r, r)

    rt_v = [dados.get("test1_rg",""), dados.get("test2_rg","")]
    def rt_r(m):
        i = rt_c[0]; rt_c[0] += 1
        return rt_v[i] if i < len(rt_v) and rt_v[i] else m.group(0)
    r = re.sub(r'\[RG\]', rt_r, r)
    return r


def substituir_documento(src, dados, dst):
    doc = Document(src)
    rg_c=[0]; cpf_c=[0]; nt_c=[0]; rt_c=[0]
    def proc(para):
        orig = para.text
        novo = substituir_paragrafo(orig, dados, rg_c, cpf_c, nt_c, rt_c)
        if novo != orig and para.runs:
            para.runs[0].text = novo
            for run in para.runs[1:]: run.text = ""
    for p in doc.paragraphs: proc(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: proc(p)
    doc.save(dst)


def docx_to_pdf(docx_path, out_dir):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        # Windows: tenta caminhos padrão
        for p in [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]:
            if os.path.exists(p): soffice = p; break
    if not soffice:
        raise FileNotFoundError(
            "LibreOffice não encontrado.\n"
            "Instale em: https://www.libreoffice.org/download/download/"
        )
    res = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        capture_output=True, text=True
    )
    if res.returncode != 0:
        raise RuntimeError(f"Erro ao converter: {res.stderr}")
    base = Path(docx_path).stem
    pdf = Path(out_dir) / f"{base}.pdf"
    if not pdf.exists():
        raise FileNotFoundError(f"PDF não gerado: {pdf}")
    return str(pdf)


def gerar_pdfs(dados, docs_sel, out_dir, callback):
    """Roda em thread separada. callback(ok, msg, arquivos)."""
    try:
        if Document is None:
            raise ImportError("python-docx não instalado. Execute: pip install python-docx")

        base_dir = resource_path("docs")
        FONTES = {
            "procuracao": os.path.join(base_dir, "PROCURACAO_MODELO.docx"),
            "contrato":   os.path.join(base_dir, "HONORARIOS_ATUALIZADO.docx"),
        }
        NOMES = {
            "procuracao": "Procuracao",
            "contrato":   "Contrato",
        }
        os.makedirs(out_dir, exist_ok=True)
        primeiro = dados.get("nome_menor","Cliente").split()[0]
        gerados = []
        for dt in docs_sel:
            src = FONTES[dt]
            if not os.path.exists(src):
                raise FileNotFoundError(f"Modelo não encontrado:\n{src}")
            nome = f"{NOMES[dt]}_{primeiro}"
            tmp  = os.path.join(out_dir, f"{nome}.docx")
            substituir_documento(src, dados, tmp)
            pdf = docx_to_pdf(tmp, out_dir)
            final = os.path.join(out_dir, f"{nome}.pdf")
            if pdf != final: shutil.move(pdf, final)
            try: os.remove(tmp)
            except: pass
            gerados.append(final)

        callback(True, f"{len(gerados)} documento(s) gerado(s) com sucesso.", gerados)
    except Exception as e:
        callback(False, str(e), [])


# ── Widgets auxiliares ───────────────────────────────────────────────────────

class GoldSeparator(tk.Canvas):
    def __init__(self, parent, **kw):
        super().__init__(parent, height=1, bg=C["bg3"],
                         highlightthickness=0, bd=0, **kw)
        self.bind("<Configure>", self._draw)

    def _draw(self, e=None):
        self.delete("all")
        w = self.winfo_width()
        self.create_line(0, 0, w, 0, fill=C["gold_dim"], width=1)


class SectionHeader(tk.Frame):
    def __init__(self, parent, text, **kw):
        super().__init__(parent, bg=C["bg3"], **kw)
        # Linha esquerda
        lf = tk.Frame(self, bg=C["gold_dim"], width=24, height=1)
        lf.pack(side="left", padx=(0,8), pady=9)
        tk.Label(self, text=text, font=FONT_SECT,
                 bg=C["bg3"], fg=C["gold"], padx=0).pack(side="left")
        # Linha direita
        rf = tk.Frame(self, bg=C["gold_dim"], height=1)
        rf.pack(side="left", fill="x", expand=True, padx=(8,0), pady=9)


class LabeledInput(tk.Frame):
    def __init__(self, parent, label, placeholder="", width=28, **kw):
        super().__init__(parent, bg=C["bg3"], **kw)
        tk.Label(self, text=label.upper(), font=FONT_LABEL,
                 bg=C["bg3"], fg=C["muted"]).pack(anchor="w", pady=(0,3))

        self.var = tk.StringVar()
        frame = tk.Frame(self, bg=C["border"], padx=1, pady=1)
        frame.pack(fill="x")
        inner = tk.Frame(frame, bg=C["accent"])
        inner.pack(fill="x")
        self.entry = tk.Entry(inner, textvariable=self.var,
                              font=FONT_INPUT, bg=C["accent"], fg=C["text"],
                              insertbackground=C["gold"],
                              relief="flat", bd=0,
                              highlightthickness=0, width=width)
        self.entry.pack(fill="x", padx=10, pady=7)

        # Placeholder
        self._ph = placeholder
        if placeholder:
            self._show_ph()
            self.entry.bind("<FocusIn>",  self._hide_ph)
            self.entry.bind("<FocusOut>", self._show_ph)

        # Hover
        for w in (frame, inner, self.entry):
            w.bind("<Enter>", lambda e: frame.config(bg=C["gold_dim"]))
            w.bind("<Leave>", lambda e: frame.config(bg=C["border"]))

    def _show_ph(self, e=None):
        if not self.var.get():
            self.entry.insert(0, self._ph)
            self.entry.config(fg=C["muted"])

    def _hide_ph(self, e=None):
        if self.entry.get() == self._ph:
            self.entry.delete(0, "end")
            self.entry.config(fg=C["text"])

    def get(self):
        v = self.var.get()
        return "" if v == self._ph else v

    def focus(self): self.entry.focus()


# ── Janela principal ─────────────────────────────────────────────────────────

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Uptech Sign — Gerador de Documentos")
        self.geometry("920x820")
        self.minsize(820, 700)
        self.configure(bg=C["bg"])
        self.resizable(True, True)

        # Ícone (ignora se não encontrar)
        try:
            self.iconbitmap(resource_path("assets/icon.ico"))
        except: pass

        self._build()
        self.center()

    def center(self):
        self.update_idletasks()
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        w, h = self.winfo_width(), self.winfo_height()
        self.geometry(f"+{(sw-w)//2}+{(sh-h)//2}")

    def _build(self):
        # ── Cabeçalho ────────────────────────────────────────────────────────
        header = tk.Frame(self, bg=C["bg2"], pady=0)
        header.pack(fill="x")

        # Barra dourada topo
        tk.Frame(header, bg=C["gold"], height=3).pack(fill="x")

        inner_h = tk.Frame(header, bg=C["bg2"])
        inner_h.pack(fill="x", padx=36, pady=20)

        # Emblema
        emblem = tk.Frame(inner_h, bg=C["bg2"])
        emblem.pack(side="left")

        em_border = tk.Frame(emblem, bg=C["gold"], padx=2, pady=2)
        em_border.pack()
        em_inner = tk.Frame(em_border, bg=C["bg2"], padx=14, pady=10)
        em_inner.pack()
        tk.Label(em_inner, text="DD", font=("Georgia", 18, "bold"),
                 bg=C["bg2"], fg=C["gold"]).pack()

        # Títulos
        titles = tk.Frame(inner_h, bg=C["bg2"])
        titles.pack(side="left", padx=20)
        tk.Label(titles, text="UPTECH SIGN",
                 font=("Georgia", 18, "bold"),
                 bg=C["bg2"], fg=C["white"]).pack(anchor="w")
        tk.Label(titles, text="Gerador de Documentos Jurídicos",
                 font=("Georgia", 10, "italic"),
                 bg=C["bg2"], fg=C["gold"]).pack(anchor="w")
        tk.Label(titles, text="Daniel Cunha Detter  ·  OAB/SP 258.095",
                 font=FONT_SMALL,
                 bg=C["bg2"], fg=C["muted"]).pack(anchor="w", pady=(4,0))

        # Linha dourada separadora
        tk.Frame(header, bg=C["gold_dim"], height=1).pack(fill="x")

        # ── Seletor de documento ──────────────────────────────────────────────
        sel_frame = tk.Frame(self, bg=C["bg"], pady=0)
        sel_frame.pack(fill="x", padx=36, pady=(20, 0))

        tk.Label(sel_frame, text="DOCUMENTO A GERAR",
                 font=FONT_LABEL, bg=C["bg"], fg=C["muted"]).pack(anchor="w", pady=(0,8))

        tabs = tk.Frame(sel_frame, bg=C["border"], padx=1, pady=1)
        tabs.pack(anchor="w")
        inner_tabs = tk.Frame(tabs, bg=C["bg3"])
        inner_tabs.pack()

        self.doc_var = tk.StringVar(value="ambos")
        self._tab_btns = {}
        options = [("Procuração + Contrato", "ambos"),
                   ("Só Procuração",         "procuracao"),
                   ("Só Contrato",           "contrato")]

        for label, val in options:
            b = tk.Button(inner_tabs, text=label,
                          font=("Georgia", 10, "bold"),
                          relief="flat", bd=0, padx=20, pady=9,
                          cursor="hand2",
                          command=lambda v=val: self._set_tab(v))
            b.pack(side="left")
            self._tab_btns[val] = b

        self._set_tab("ambos")

        # ── Scroll frame principal ────────────────────────────────────────────
        outer = tk.Frame(self, bg=C["bg"])
        outer.pack(fill="both", expand=True, padx=36, pady=16)

        canvas = tk.Canvas(outer, bg=C["bg"], highlightthickness=0, bd=0)
        scrollbar = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        self.scroll_frame = tk.Frame(canvas, bg=C["bg3"],
                                     highlightthickness=1,
                                     highlightbackground=C["border"])
        self.sf_id = canvas.create_window((0,0), window=self.scroll_frame, anchor="nw")

        def on_resize(e):
            canvas.itemconfig(self.sf_id, width=e.width)
        canvas.bind("<Configure>", on_resize)
        self.scroll_frame.bind("<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        def _scroll(e):
            canvas.yview_scroll(-1*(e.delta//120), "units")
        canvas.bind_all("<MouseWheel>", _scroll)

        self._build_form(self.scroll_frame)

        # ── Rodapé ────────────────────────────────────────────────────────────
        footer = tk.Frame(self, bg=C["bg2"])
        footer.pack(fill="x", side="bottom")
        tk.Frame(footer, bg=C["gold_dim"], height=1).pack(fill="x")
        foot_inner = tk.Frame(footer, bg=C["bg2"])
        foot_inner.pack(fill="x", padx=36, pady=12)

        self.status_var = tk.StringVar(value="Pronto para gerar documentos.")
        self.status_lbl = tk.Label(foot_inner, textvariable=self.status_var,
                                   font=FONT_STATUS, bg=C["bg2"], fg=C["muted"],
                                   anchor="w")
        self.status_lbl.pack(side="left", fill="x", expand=True)

        tk.Label(foot_inner, text="OAB/SP 258.095",
                 font=FONT_SMALL, bg=C["bg2"], fg=C["gold_dim"]).pack(side="right")

    def _set_tab(self, val):
        self.doc_var.set(val)
        for k, b in self._tab_btns.items():
            if k == val:
                b.config(bg=C["gold"], fg=C["bg2"],
                         activebackground=C["gold_light"], activeforeground=C["bg2"])
            else:
                b.config(bg=C["bg3"], fg=C["muted"],
                         activebackground=C["accent"], activeforeground=C["text"])

        # Mostrar/ocultar campos exclusivos
        show_proc    = val in ("ambos", "procuracao")
        show_contrato = val in ("ambos", "contrato")

        for w in getattr(self, "_proc_only", []):
            w.pack_info_stored = getattr(w, "_pack_info", None)
            if show_proc:
                w.pack(fill="x", pady=4)
            else:
                w.pack_forget()

        for w in getattr(self, "_cont_only", []):
            if show_contrato:
                w.pack(fill="x", pady=4)
            else:
                w.pack_forget()

    def _build_form(self, parent):
        pad = {"padx": 28, "pady": 0}
        grid_pad = {"padx": 28}

        def section(text):
            f = tk.Frame(parent, bg=C["bg3"])
            f.pack(fill="x", padx=28, pady=(18, 4))
            tk.Frame(f, bg=C["gold"], width=4, height=22).pack(side="left", padx=(0,10))
            tk.Label(f, text=text, font=FONT_SECT,
                     bg=C["bg3"], fg=C["gold"]).pack(side="left")
            tk.Frame(f, bg=C["gold_dim"], height=1).pack(
                side="left", fill="x", expand=True, padx=(12,0), pady=10)

        def row2(f1, f2):
            r = tk.Frame(parent, bg=C["bg3"])
            r.pack(fill="x", **grid_pad, pady=4)
            f1.pack(side="left", fill="x", expand=True, in_=r, padx=(0,8))
            f2.pack(side="left", fill="x", expand=True, in_=r)
            return r

        def field(label, ph="", full=False):
            f = LabeledInput(parent, label, placeholder=ph)
            if full: f.pack(fill="x", **grid_pad, pady=4)
            return f

        # ── Menor ────────────────────────────────────────────────────────────
        section("MENOR / INCAPAZ")
        self.f_nome_menor = field("Nome do Menor / Incapaz",
                                  "Nome completo", full=True)
        r1a = tk.Frame(parent, bg=C["bg3"])
        r1a.pack(fill="x", **grid_pad, pady=4)
        self.f_nasc  = LabeledInput(r1a, "Data de Nascimento", "DD/MM/AAAA")
        self.f_rg_m  = LabeledInput(r1a, "RG do Menor", "00.000.000-0")
        self.f_cpf_m = LabeledInput(r1a, "CPF do Menor", "000.000.000-00")
        self.f_nasc.pack(side="left", fill="x", expand=True, padx=(0,6))
        self.f_rg_m.pack(side="left", fill="x", expand=True, padx=(0,6))
        self.f_cpf_m.pack(side="left", fill="x", expand=True)

        # ── Representante ────────────────────────────────────────────────────
        section("REPRESENTANTE LEGAL")
        self.f_nome_rep = field("Nome da Representante Legal",
                                "Nome completo", full=True)
        r2a = tk.Frame(parent, bg=C["bg3"])
        r2a.pack(fill="x", **grid_pad, pady=4)
        self.f_estcivil = LabeledInput(r2a, "Estado Civil", "ex: solteira")
        self.f_prof     = LabeledInput(r2a, "Profissão",    "ex: comerciante")
        self.f_estcivil.pack(side="left", fill="x", expand=True, padx=(0,8))
        self.f_prof.pack(side="left", fill="x", expand=True)

        r2b = tk.Frame(parent, bg=C["bg3"])
        r2b.pack(fill="x", **grid_pad, pady=4)
        self.f_rg_r  = LabeledInput(r2b, "RG da Representante", "00.000.000-0")
        self.f_cpf_r = LabeledInput(r2b, "CPF da Representante", "000.000.000-00")
        self.f_rg_r.pack(side="left", fill="x", expand=True, padx=(0,8))
        self.f_cpf_r.pack(side="left", fill="x", expand=True)

        self.f_end = field("Endereço Completo com CEP",
                           "Rua, nº, bairro, cidade – SP, CEP 00000-000", full=True)

        # ── Seção Procuração (e-mail) ─────────────────────────────────────────
        self._proc_only = []
        sec_proc = tk.Frame(parent, bg=C["bg3"])
        sec_proc.pack(fill="x", padx=28, pady=(18,4))
        tk.Frame(sec_proc, bg=C["gold"], width=4, height=22).pack(side="left", padx=(0,10))
        tk.Label(sec_proc, text="PROCURAÇÃO", font=FONT_SECT,
                 bg=C["bg3"], fg=C["gold"]).pack(side="left")
        tk.Frame(sec_proc, bg=C["gold_dim"], height=1).pack(
            side="left", fill="x", expand=True, padx=(12,0), pady=10)
        self._proc_only.append(sec_proc)

        self.f_email = field("E-mail do Advogado", "contato@detter.adv.br")
        self.f_email.pack(fill="x", **grid_pad, pady=4)
        self._proc_only.append(self.f_email)

        # ── Seção Contrato (data + testemunhas) ───────────────────────────────
        self._cont_only = []
        sec_cont = tk.Frame(parent, bg=C["bg3"])
        sec_cont.pack(fill="x", padx=28, pady=(18,4))
        tk.Frame(sec_cont, bg=C["gold"], width=4, height=22).pack(side="left", padx=(0,10))
        tk.Label(sec_cont, text="CONTRATO — DATA & ASSINATURAS", font=FONT_SECT,
                 bg=C["bg3"], fg=C["gold"]).pack(side="left")
        tk.Frame(sec_cont, bg=C["gold_dim"], height=1).pack(
            side="left", fill="x", expand=True, padx=(12,0), pady=10)
        self._cont_only.append(sec_cont)

        r_data = tk.Frame(parent, bg=C["bg3"])
        r_data.pack(fill="x", **grid_pad, pady=4)
        self.f_cidade = LabeledInput(r_data, "Cidade",    "ex: Santos")
        self.f_dia    = LabeledInput(r_data, "Dia",       "ex: 18")
        self.f_mes    = LabeledInput(r_data, "Mês",       "ex: março")
        self.f_ano    = LabeledInput(r_data, "Ano",       "ex: 2026")
        self.f_cidade.pack(side="left", fill="x", expand=True, padx=(0,6))
        self.f_dia.pack(side="left", fill="x", expand=True, padx=(0,6))
        self.f_mes.pack(side="left", fill="x", expand=True, padx=(0,6))
        self.f_ano.pack(side="left", fill="x", expand=True)
        self._cont_only.append(r_data)

        r_t1 = tk.Frame(parent, bg=C["bg3"])
        r_t1.pack(fill="x", **grid_pad, pady=4)
        self.f_t1n = LabeledInput(r_t1, "Testemunha 1 — Nome",  "Nome completo")
        self.f_t1r = LabeledInput(r_t1, "Testemunha 1 — RG",    "00.000.000-0")
        self.f_t1n.pack(side="left", fill="x", expand=True, padx=(0,8))
        self.f_t1r.pack(side="left", fill="x", expand=True)
        self._cont_only.append(r_t1)

        r_t2 = tk.Frame(parent, bg=C["bg3"])
        r_t2.pack(fill="x", **grid_pad, pady=4)
        self.f_t2n = LabeledInput(r_t2, "Testemunha 2 — Nome",  "Nome completo")
        self.f_t2r = LabeledInput(r_t2, "Testemunha 2 — RG",    "00.000.000-0")
        self.f_t2n.pack(side="left", fill="x", expand=True, padx=(0,8))
        self.f_t2r.pack(side="left", fill="x", expand=True)
        self._cont_only.append(r_t2)

        # ── Pasta de saída ────────────────────────────────────────────────────
        section("SALVAR EM")
        out_row = tk.Frame(parent, bg=C["bg3"])
        out_row.pack(fill="x", **grid_pad, pady=(4,0))

        out_border = tk.Frame(out_row, bg=C["border"], padx=1, pady=1)
        out_border.pack(side="left", fill="x", expand=True)
        out_inner = tk.Frame(out_border, bg=C["accent"])
        out_inner.pack(fill="x")
        self.out_var = tk.StringVar(value=os.path.join(
            os.path.expanduser("~"), "Desktop", "Documentos_Detter"))
        tk.Entry(out_inner, textvariable=self.out_var,
                 font=FONT_INPUT, bg=C["accent"], fg=C["text"],
                 insertbackground=C["gold"],
                 relief="flat", bd=0, highlightthickness=0).pack(
            fill="x", padx=10, pady=7)

        browse_btn = tk.Button(out_row, text="  📁  ",
                               font=("Georgia", 11),
                               bg=C["accent"], fg=C["gold"],
                               activebackground=C["gold"], activeforeground=C["bg"],
                               relief="flat", bd=0, padx=12, pady=6, cursor="hand2",
                               command=self._browse)
        browse_btn.pack(side="left", padx=(8,0))

        # ── Botão Gerar ───────────────────────────────────────────────────────
        btn_frame = tk.Frame(parent, bg=C["bg3"])
        btn_frame.pack(fill="x", padx=28, pady=24)

        self.gen_btn = tk.Button(
            btn_frame,
            text="⚖  GERAR DOCUMENTOS  ⚖",
            font=("Georgia", 13, "bold"),
            bg=C["gold"], fg=C["bg"],
            activebackground=C["gold_light"], activeforeground=C["bg"],
            relief="flat", bd=0,
            padx=32, pady=14,
            cursor="hand2",
            command=self._gerar
        )
        self.gen_btn.pack(fill="x")

        # Resultado
        self.result_frame = tk.Frame(parent, bg=C["bg3"])
        self.result_frame.pack(fill="x", padx=28, pady=(0,28))

        # Inicializar visibilidade das abas
        self._set_tab("ambos")

    def _browse(self):
        d = filedialog.askdirectory(title="Escolher pasta de destino")
        if d: self.out_var.set(d)

    def _collect(self):
        return {
            "nome_menor":   self.f_nome_menor.get(),
            "nascimento":   self.f_nasc.get(),
            "rg_menor":     self.f_rg_m.get(),
            "cpf_menor":    self.f_cpf_m.get(),
            "nome_rep":     self.f_nome_rep.get(),
            "estado_civil": self.f_estcivil.get(),
            "profissao":    self.f_prof.get(),
            "rg_rep":       self.f_rg_r.get(),
            "cpf_rep":      self.f_cpf_r.get(),
            "endereco":     self.f_end.get(),
            "email":        self.f_email.get(),
            "cidade":       self.f_cidade.get(),
            "data_dia":     self.f_dia.get(),
            "mes":          self.f_mes.get(),
            "ano":          self.f_ano.get(),
            "test1_nome":   self.f_t1n.get(),
            "test1_rg":     self.f_t1r.get(),
            "test2_nome":   self.f_t2n.get(),
            "test2_rg":     self.f_t2r.get(),
        }

    def _gerar(self):
        dados = self._collect()
        if not dados["nome_menor"]:
            messagebox.showwarning("Campo obrigatório",
                                   "Informe o nome do menor / incapaz.")
            return

        tab = self.doc_var.get()
        docs_sel = (["procuracao", "contrato"] if tab == "ambos"
                    else [tab])
        out_dir = self.out_var.get() or os.path.join(
            os.path.expanduser("~"), "Desktop", "Documentos_Detter")

        # Limpar resultado anterior
        for w in self.result_frame.winfo_children(): w.destroy()

        self.gen_btn.config(state="disabled", text="Gerando…  ⌛")
        self._set_status("Processando documentos…", C["gold"])

        def callback(ok, msg, arquivos):
            self.after(0, lambda: self._on_done(ok, msg, arquivos))

        threading.Thread(
            target=gerar_pdfs,
            args=(dados, docs_sel, out_dir, callback),
            daemon=True
        ).start()

    def _on_done(self, ok, msg, arquivos):
        self.gen_btn.config(state="normal",
                            text="⚖  GERAR DOCUMENTOS  ⚖")
        if ok:
            self._set_status(f"✓  {msg}", C["success_fg"])
            self._show_results(arquivos)
        else:
            self._set_status(f"✗  Erro ao gerar", C["error_fg"])
            messagebox.showerror("Erro na geração", msg)

    def _show_results(self, arquivos):
        for w in self.result_frame.winfo_children(): w.destroy()

        # Caixa de sucesso
        box = tk.Frame(self.result_frame, bg=C["success"], padx=1, pady=1)
        box.pack(fill="x", pady=(8,0))
        inner = tk.Frame(box, bg=C["bg3"])
        inner.pack(fill="x")
        tk.Label(inner, text="✓  Documentos gerados com sucesso",
                 font=("Georgia", 11, "bold"),
                 bg=C["bg3"], fg=C["success_fg"],
                 pady=10).pack(anchor="w", padx=16)

        for arq in arquivos:
            row = tk.Frame(inner, bg=C["bg3"])
            row.pack(fill="x", padx=16, pady=(0,6))
            nome = os.path.basename(arq)
            tk.Label(row, text=f"📄  {nome}", font=FONT_STATUS,
                     bg=C["bg3"], fg=C["text"]).pack(side="left")
            tk.Button(row, text="Abrir",
                      font=("Georgia", 9, "bold"),
                      bg=C["gold_dim"], fg=C["text"],
                      activebackground=C["gold"], activeforeground=C["bg"],
                      relief="flat", bd=0, padx=12, pady=3,
                      cursor="hand2",
                      command=lambda a=arq: self._abrir(a)).pack(side="right", padx=(0,0))

        tk.Button(inner, text="  📁  Abrir pasta",
                  font=("Georgia", 10, "bold"),
                  bg=C["accent"], fg=C["gold"],
                  activebackground=C["gold"], activeforeground=C["bg"],
                  relief="flat", bd=0, padx=14, pady=7,
                  cursor="hand2",
                  command=lambda: self._abrir_pasta(os.path.dirname(arquivos[0]))
                  ).pack(anchor="w", padx=16, pady=(2,12))

    def _abrir(self, path):
        try:
            if sys.platform == "win32": os.startfile(path)
            elif sys.platform == "darwin": subprocess.run(["open", path])
            else: subprocess.run(["xdg-open", path])
        except Exception as e:
            messagebox.showerror("Erro", str(e))

    def _abrir_pasta(self, path):
        try:
            if sys.platform == "win32":
                subprocess.run(["explorer", path])
            elif sys.platform == "darwin":
                subprocess.run(["open", path])
            else:
                subprocess.run(["xdg-open", path])
        except: pass

    def _set_status(self, msg, color=None):
        self.status_var.set(msg)
        if color: self.status_lbl.config(fg=color)


if __name__ == "__main__":
    app = App()
    app.mainloop()
