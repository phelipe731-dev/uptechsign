"""Document generation service — ported from app.py lines 61-184."""

import asyncio
import os
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document

from app.config import settings
from app.utils.file_storage import document_dir, random_filename
from app.utils.security import sha256_file


def substituir_paragrafo(text: str, dados: dict, rg_c: list, cpf_c: list, nt_c: list, rt_c: list) -> str:
    """Replace placeholder fields in a paragraph's text."""
    r = text
    simples = {
        "NOME DO MENOR/INCAPAZ":      dados.get("nome_menor", ""),
        "Data de Nascimento":          dados.get("nascimento", ""),
        "NOME DA REPRESENTANTE LEGAL": dados.get("nome_rep", ""),
        "estado civil":                dados.get("estado_civil", ""),
        "profissão":                   dados.get("profissao", ""),
        "Endereço completo com CEP":   dados.get("endereco", ""),
        "Cidade":                      dados.get("cidade", ""),
        "data":                        dados.get("data_dia", ""),
        "mês":                         dados.get("mes", ""),
        "ano":                         dados.get("ano", ""),
        "e-mail do advogado":          dados.get("email", ""),
    }
    for campo, val in simples.items():
        if val:
            r = r.replace(f"[{campo}]", val)

    rg_v = [dados.get("rg_menor", ""), dados.get("rg_rep", "")]
    def rg_r(m):
        i = rg_c[0]; rg_c[0] += 1
        return rg_v[i] if i < len(rg_v) and rg_v[i] else m.group(0)
    r = re.sub(r'\[nº do RG\]', rg_r, r)

    cpf_v = [dados.get("cpf_menor", ""), dados.get("cpf_rep", "")]
    def cpf_r(m):
        i = cpf_c[0]; cpf_c[0] += 1
        return cpf_v[i] if i < len(cpf_v) and cpf_v[i] else m.group(0)
    r = re.sub(r'\[nº do CPF\]', cpf_r, r)

    nt_v = [dados.get("test1_nome", ""), dados.get("test2_nome", "")]
    def nt_r(m):
        i = nt_c[0]; nt_c[0] += 1
        return nt_v[i] if i < len(nt_v) and nt_v[i] else m.group(0)
    r = re.sub(r'\[NOME\]', nt_r, r)

    rt_v = [dados.get("test1_rg", ""), dados.get("test2_rg", "")]
    def rt_r(m):
        i = rt_c[0]; rt_c[0] += 1
        return rt_v[i] if i < len(rt_v) and rt_v[i] else m.group(0)
    r = re.sub(r'\[RG\]', rt_r, r)

    return r


def substituir_documento(src: str, dados: dict, dst: str) -> None:
    """Apply field substitutions to a DOCX template and save to dst."""
    doc = Document(src)
    rg_c = [0]; cpf_c = [0]; nt_c = [0]; rt_c = [0]

    def proc(para):
        orig = para.text
        novo = substituir_paragrafo(orig, dados, rg_c, cpf_c, nt_c, rt_c)
        if novo != orig and para.runs:
            para.runs[0].text = novo
            for run in para.runs[1:]:
                run.text = ""

    for p in doc.paragraphs:
        proc(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    proc(p)
    doc.save(dst)


def docx_to_pdf_sync(docx_path: str, out_dir: str) -> str:
    """Convert DOCX to PDF using LibreOffice. Returns the PDF path."""
    soffice = settings.LIBREOFFICE_PATH
    if soffice == "soffice":
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            for p in [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]:
                if os.path.exists(p):
                    soffice = p
                    break
    if not soffice:
        raise FileNotFoundError(
            "LibreOffice não encontrado. "
            "Instale em: https://www.libreoffice.org/download/download/"
        )
    res = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        capture_output=True, text=True,
    )
    if res.returncode != 0:
        raise RuntimeError(f"Erro ao converter: {res.stderr}")
    base = Path(docx_path).stem
    pdf = Path(out_dir) / f"{base}.pdf"
    if not pdf.exists():
        raise FileNotFoundError(f"PDF não gerado: {pdf}")
    return str(pdf)


async def docx_to_pdf(docx_path: str, out_dir: str) -> str:
    """Async wrapper for DOCX to PDF conversion."""
    return await asyncio.to_thread(docx_to_pdf_sync, docx_path, out_dir)


def parse_template_placeholders(docx_path: str) -> list[str]:
    """Scan a DOCX template and return unique [CAMPO] placeholder labels found."""
    doc = Document(docx_path)
    found: list[str] = []
    seen: set[str] = set()

    def scan_text(text: str) -> None:
        for m in re.finditer(r'\[([^\[\]]+)\]', text):
            label = m.group(1)
            if label not in seen:
                seen.add(label)
                found.append(label)

    for p in doc.paragraphs:
        scan_text(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan_text(p.text)

    return found


def label_to_key(label: str) -> str:
    """Convert a placeholder label to a snake_case field key."""
    key = label.lower()
    key = re.sub(r'[^\w\s]', '', key)
    key = re.sub(r'\s+', '_', key.strip())
    key = re.sub(r'_+', '_', key)
    return key or "field"


def substitute_generic_sync(
    src: str,
    field_data: dict,
    template_fields: list[dict],
    dst: str,
) -> None:
    """Generic substitution using template fields [{key, label, ...}].

    Supports positional duplicates: if multiple fields share the same label
    (e.g. two fields with label 'nº do RG'), each occurrence in the document
    is replaced by the corresponding field value in display_order sequence.
    """
    from collections import defaultdict

    doc = Document(src)

    # Build label → [value1, value2, ...] in display_order
    label_values: dict[str, list[str]] = defaultdict(list)
    for f in sorted(template_fields, key=lambda x: x.get("display_order", 0)):
        label_values[f["label"]].append(str(field_data.get(f["key"], "")))

    # Document-wide occurrence counter per label
    occurrence_counter: dict[str, int] = defaultdict(int)

    def replace_in_text(text: str) -> str:
        def replacer(m: re.Match) -> str:
            label = m.group(1)
            if label not in label_values:
                return m.group(0)
            values = label_values[label]
            idx = occurrence_counter[label]
            occurrence_counter[label] += 1
            if idx < len(values) and values[idx]:
                return values[idx]
            return m.group(0)
        return re.sub(r'\[([^\[\]]+)\]', replacer, text)

    def proc(para) -> None:
        orig = para.text
        novo = replace_in_text(orig)
        if novo != orig and para.runs:
            para.runs[0].text = novo
            for run in para.runs[1:]:
                run.text = ""

    for p in doc.paragraphs:
        proc(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    proc(p)

    doc.save(dst)


async def generate_document(
    template_path: str,
    field_data: dict,
    document_id: str | None,
    template_fields: list[dict] | None = None,
    output_dir: str | Path | None = None,
) -> tuple[str, str, str, str]:
    """
    Generate a document from a template.
    If template_fields is provided, uses generic label-based substitution.
    Otherwise falls back to the legacy hardcoded substitution.
    Returns (docx_path, pdf_path, docx_sha256, pdf_sha256).
    """
    if output_dir is not None:
        doc_dir = Path(output_dir)
        doc_dir.mkdir(parents=True, exist_ok=True)
    elif document_id:
        doc_dir = document_dir(document_id)
    else:
        raise ValueError("document_id ou output_dir deve ser informado para gerar o documento.")
    docx_name = random_filename(".docx")
    pdf_name = random_filename(".pdf")

    docx_path = str(doc_dir / docx_name)

    # Run template substitution in thread (CPU-bound)
    if template_fields:
        await asyncio.to_thread(substitute_generic_sync, template_path, field_data, template_fields, docx_path)
    else:
        await asyncio.to_thread(substituir_documento, template_path, field_data, docx_path)

    # Convert to PDF
    temp_pdf = await docx_to_pdf(docx_path, str(doc_dir))

    # Rename to random name
    pdf_path = str(doc_dir / pdf_name)
    if temp_pdf != pdf_path:
        shutil.move(temp_pdf, pdf_path)

    docx_hash = sha256_file(docx_path)
    pdf_hash = sha256_file(pdf_path)

    return docx_path, pdf_path, docx_hash, pdf_hash
